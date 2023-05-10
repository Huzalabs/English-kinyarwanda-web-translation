import streamlit as st 
from translator import *
import pandas as pd
from accuracy import *
from docx import Document
from annotated_text import annotated_text


st.title("GOV Contents Translation")

col1, col2 = st.columns(2)

with col1:
   st.image("https://www.developmentaid.org/files/organizationLogos/huzalabs-435178.jpg")

with col2:
   st.image("/Users/kayarn/Desktop/Sources/NLP_Projects/RW-EN-Gov-Web-Translation/images/Rwanda Gov.png")

st.image("/Users/kayarn/Desktop/Sources/NLP_Projects/RW-EN-Gov-Web-Translation/images/GIZ_HUZALABS_FAIRFORWARD.png")

# with st.sidebar:
#     st.header("Choose an option")
# Add CSS styling to center the element
st.markdown(
    """
    <style>
    .center {
        display: flex;
        justify-content: center;
        align-items: center;
        height: 80px;
        margin: 0 auto;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

# Add the file uploader element inside a div with the center class
with st.container():
    st.header("Upload your Docx file")
    file_upload = st.file_uploader('', type=['docx'])
    

    #save file file
    if file_upload:
        with open('data/content.docx', "wb+") as doc_file:
            doc_file.write(file_upload.read())
        #Read the doc file
        doc = Document('data/content.docx')
        paragraphs = []
        for i in range(len(doc.paragraphs)):
            paragraphs.append(doc.paragraphs[i].text)
        st.session_state['text_area'] = "\n".join(paragraphs)

    st.header("Or Paste/Type your Text Here")
    #Paste in your text docs
    st.text_area("", key='text_area', height=10)


    
    
translate_button, download_button, a ,b  = st.columns(4)

with translate_button:
    translate_button = st.button("Translate Dataset")
   
       
if translate_button:
    paragraphs = st.session_state['text_area']
    print(paragraphs)
    texts = []
    for text in paragraphs:
        trans_text = translate_row(paragraphs, target_lang='rw', source_lang='en')
        texts.append(trans_text)

    #st.text_area("", key="translate_area", disabled=True, height=20)


with download_button:
    download_button = st.button("Download Dataset")




# def process_upload(df):
#     with st.spinner("Translating ......."):
#         df['translated_text']=translate_dataframe(df, target_lang='rw', source_lang='en')
#         df['Back_translated_text']=back_translate_dataframe(df, target_lang='en', source_lang='rw')

#     return df

# def calc_accuracy(df):
#     # Iterate over the DataFrame rows and compare the paragraphs in each row
#     for i, row in df.iterrows():
#         similarity = compare_paragraphs(row['content'], row['Back_translated_text'])
#         df.at[i, 'Condidence level'] = similarity
#     return df[['content','translated_text','Condidence level']]


# # Print the updated dataframe
# if file_upload:
#     if st.button('Translate'):
#         #Save the csv file
#         with open('data/dataset.csv', 'wb+') as file:
#             file.write(file_upload.read())
        
#         df = pd.read_csv('data/dataset.csv')
#         df = process_upload(df)

#         st.header("Displaying the confidence level of the translation")
#         df = calc_accuracy(df)

#         st.dataframe(df)

#         #Save the output
#         df.to_csv('output.csv')


#         def convert_df(df):
#             # IMPORTANT: Cache the conversion to prevent computation on every rerun
#             return df.to_csv().encode('utf-8')

#         csv = convert_df(df)

#         st.download_button(label="Download a CSV",file_name='output.csv',data=csv, mime='text/csv')





        



